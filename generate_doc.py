from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_doc():
    doc = Document()

    # Title
    title = doc.add_heading('OnGround AI Assistant', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Subtitle
    subtitle = doc.add_paragraph('Functional Overview & User Documentation')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # 1. Introduction
    doc.add_heading('1. Executive Summary', level=1)
    doc.add_paragraph(
        "OnGround AI is a strategic intelligence platform designed to empower Area Sales Officers and "
        "Retail Executives with data-driven insights. By combining real-time market metrics with "
        "advanced Generative AI, the platform transforms raw field data into actionable business strategies."
    )

    # 2. Key Features
    doc.add_heading('2. Key Features', level=1)
    
    doc.add_heading('2.1 Market Intelligence Dashboard', level=2)
    doc.add_paragraph(
        "Provides a 360-degree view of market performance. Key components include:"
    )
    list_items = [
        "KPI Cards: Instant visibility into Outlet Counts, Revenue, and Growth Index.",
        "Interactive Mapping: Visualize store clusters and zoom into specific markets like Cyber City.",
        "Sales Trends: Month-over-Month performance tracking with automated growth analysis."
    ]
    for item in list_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('2.2 OnGround AI Assistant', level=2)
    doc.add_paragraph(
        "A ChatGPT-like interface tailored for retail. Features include:"
    )
    list_items = [
        "Role-Based Insights: Interact with the AI as an Area Sales Officer or Analyst.",
        "Context Awareness: The AI analyzes your current dashboard data to answer specific questions.",
        "Multi-Modal Analysis: Upload photos of store shelves or inventory logs for instant AI auditing."
    ]
    for item in list_items:
        doc.add_paragraph(item, style='List Bullet')

    # 3. User Guide
    doc.add_heading('3. User Guide', level=1)
    
    doc.add_paragraph("Step 1: Navigate the Dashboard", style='Heading 2')
    doc.add_paragraph(
        "Use the sidebar to switch between the Dashboard and the AI Assistant. On the Dashboard, "
        "you can filter by specific Area Sales Officers (ASOs) to see regional data."
    )

    doc.add_paragraph("Step 2: Interactive Map Navigation", style='Heading 2')
    doc.add_paragraph(
        "Click on market region buttons (e.g., 'Cyber City 8') to automatically center the map on "
        "that territory. You can also click on Top Stores to see their exact geographic location."
    )

    doc.add_paragraph("Step 3: Chat with the Assistant", style='Heading 2')
    doc.add_paragraph(
        "In the AI Assistant tab, use the '+' button to upload images or documents. "
        "You can use the 'History' sidebar to revisit previous conversations."
    )

    # 4. Data Sources & Accuracy
    doc.add_heading('4. Data Sources & Accuracy', level=1)
    doc.add_paragraph(
        "All intelligence is derived from the master retail database. While the AI is highly "
        "capable, it is recommended to verify critical financial metrics with official "
        "market documentation provided in the system."
    )

    # Footer
    doc.add_paragraph('\n\n--- End of Document ---').alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save('OnGround_AI_Functional_Documentation.docx')
    print("Document created successfully: OnGround_AI_Functional_Documentation.docx")

if __name__ == "__main__":
    create_doc()
